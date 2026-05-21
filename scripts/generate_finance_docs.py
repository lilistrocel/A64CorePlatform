"""
Generate two finance-module .docx deliverables:
  1. FINANCE_PM_REPORT.docx  — project-manager status report
  2. FINANCE_USER_GUIDE.docx — end-user tutorial

Both files land in Docs/4-Finance-Mod-docs/.
"""

from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt, RGBColor, Cm

OUT_DIR = Path(__file__).resolve().parent.parent / "Docs" / "4-Finance-Mod-docs"
OUT_DIR.mkdir(parents=True, exist_ok=True)


# ──────────────────────────────────────────────────────────────────────────────
# Common formatting helpers
# ──────────────────────────────────────────────────────────────────────────────

def style_doc(doc: Document) -> None:
    """Apply a consistent base style to the document."""
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    # Reduce default paragraph spacing
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25


def add_h(doc: Document, text: str, level: int = 1) -> None:
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x36, 0x56)


def add_p(doc: Document, text: str, bold: bool = False, italic: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic


def add_bullet(doc: Document, text: str, level: int = 0) -> None:
    p = doc.add_paragraph(text, style="List Bullet")
    p.paragraph_format.left_indent = Cm(0.75 + level * 0.6)


def add_code(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(10)
    p.paragraph_format.left_indent = Cm(0.5)


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for run in hdr_cells[i].paragraphs[0].runs:
            run.bold = True
    for r_idx, row in enumerate(rows, start=1):
        for c_idx, val in enumerate(row):
            table.rows[r_idx].cells[c_idx].text = val
            for p in table.rows[r_idx].cells[c_idx].paragraphs:
                p.paragraph_format.space_after = Pt(2)
    # spacer after
    doc.add_paragraph()


# ──────────────────────────────────────────────────────────────────────────────
# 1. PROJECT MANAGER REPORT
# ──────────────────────────────────────────────────────────────────────────────

def build_pm_report() -> Document:
    doc = Document()
    style_doc(doc)

    # Title
    title = doc.add_heading("A64 Core Platform — Finance Module Status Report", level=0)
    for run in title.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x36, 0x56)
    sub = doc.add_paragraph()
    sub_run = sub.add_run("Prepared for the project manager")
    sub_run.italic = True
    sub_run.font.size = Pt(11)
    meta = doc.add_paragraph()
    meta.add_run("Date: 2026-05-20").italic = True
    meta.add_run("    ·    Author: Viet Anh").italic = True

    # Executive summary
    add_h(doc, "Executive Summary", 1)
    add_p(doc,
        "Construction of the A64 Core Platform finance module is approximately "
        "40% complete on the procure-to-pay (P2P) cycle. The accounting foundation "
        "(chart of accounts, approval rules, posting defaults, journal entry "
        "tables) and the first half of the transactional cycle (purchase request, "
        "purchase order, goods receipt with automatic journal entry posting) are "
        "live and demonstrable end-to-end in the development environment. "
        "Vendor invoices (with three-way match) and payments remain to be built — "
        "those modules are scoped and sequenced.")
    add_p(doc,
        "The module operates as a separate finance service (FastAPI + MySQL) that "
        "communicates with the main operational service (FastAPI + MongoDB) "
        "through a transactional outbox pattern. Mongo was migrated to a single-"
        "node replica set to make these writes atomic. Customers without finance "
        "do not deploy the finance service — operational features continue to work "
        "without it.")

    # What's working today
    add_h(doc, "What is Working Today", 1)
    add_p(doc, "The following modules and pages are live in the development environment "
              "and have been verified end-to-end through a real journal-entry post.")
    add_table(doc,
        headers=["Capability", "Status", "Brief description"],
        rows=[
            ["Chart of Accounts", "Live",
             "227 GL accounts seeded across 9 IFRS-aligned drawers. Hierarchy view, "
             "filtering, create/edit/deactivate. Supports the account_level "
             "(drawer/title/active) and account_role classifications used by the "
             "posting engine."],
            ["Approval Rules", "Live",
             "Configurable per-doc-type approval thresholds with role assignments. "
             "Already wired to the operational PR/PO submission flow; the engine "
             "queries finance to resolve whether each document needs approval."],
            ["Posting Setup", "Live",
             "Controller-facing page where each company maps system events "
             "(AP control, bank, GR/IR clearing, input/output VAT, retained "
             "earnings, etc.) to specific GL accounts. Without this, the posting "
             "engine cannot run."],
            ["Item GL Mapping", "Live",
             "Finance-side page where each purchase item is assigned its specific "
             "Inventory and Cost-of-Sales GL accounts. Auto-defaults applied on "
             "item creation, with manual override per item."],
            ["Journal Entries", "Live",
             "Read-only register of every posted JE. Row-expand shows debit and "
             "credit lines with account names. Cross-linked from the source "
             "operational document."],
            ["Goods Receipts (operational)", "Live",
             "Warehouse confirms physical receipt from an open Purchase Order. "
             "Posting the GR triggers the first real journal entry — debit "
             "inventory, credit GR/IR clearing — in the finance ledger."],
            ["P&L Statement", "Live (existing)",
             "Earlier-built P&L page integrated into the new Finance sidebar group."],
            ["Incoming Preview", "Live",
             "Finance-side read-through view of operational documents currently "
             "in Pending Approval, so finance has visibility into what is about "
             "to materialise."],
            ["Approval Inbox (purchasing)", "Live",
             "Reviewers see documents awaiting their approval and can decide "
             "without leaving the finance workflow."],
        ])

    # Architecture
    add_h(doc, "Architecture in Brief", 1)
    add_bullet(doc,
        "Two-service architecture. The main operational service owns workflow "
        "documents (PR, PO, GR). The finance service owns accounting documents "
        "(JE, AP, Payment) and runs on a separate database (MySQL).")
    add_bullet(doc,
        "Cross-service communication is one-directional via an outbox pattern. "
        "When an operational document reaches a finance-relevant state, the main "
        "service writes an event row to a shared queue. A worker drains the queue "
        "into the finance service.")
    add_bullet(doc,
        "All writes are transactional. The operational document update and the "
        "outbox event are written atomically (Mongo replica set transactions), "
        "so a failed outbox write rolls back the document update. A periodic "
        "sweeper provides defence-in-depth against any missed events.")
    add_bullet(doc,
        "The finance module is opt-in per customer. Customers who do not need "
        "accounting deploy only the operational service.")
    add_bullet(doc,
        "Authentication is shared. Users log in once; the same JWT is accepted "
        "by both services.")

    # Build phases
    add_h(doc, "Build Phases and Current Position", 1)
    add_table(doc,
        headers=["Phase", "Scope", "Status"],
        rows=[
            ["Phase 1A (master data)",
             "Vendors, items, payment terms, tax codes, fiscal periods, companies, "
             "cost centres, chart of accounts.", "Complete"],
            ["Phase 1B (purchasing)",
             "Purchase Requests, Purchase Orders, multi-step approval flow.",
             "Complete"],
            ["Phase 2 (drift mitigation)",
             "Transactional outbox with single-node Mongo replica set; "
             "periodic outbox reconciliation sweeper.", "Complete"],
            ["Phase A (foundation)",
             "Journal entry tables, posting setup, posting setup UI, item GL "
             "mapping.", "Complete"],
            ["Phase B (first posting)",
             "Goods Receipt module, purchase_received handler, JE list UI. "
             "Produces the first real accounting entry.", "Complete"],
            ["Phase C (vendor invoice)",
             "AP Invoice module on operation side; three-way match logic; "
             "ap_invoice_posted handler that recognises VAT and creates the "
             "vendor liability.", "Not started"],
            ["Phase D (payment)",
             "Payment recording on finance side; vendor_payment handler that "
             "clears AP and credits the bank.", "Not started"],
            ["Phase E (polish)",
             "GR/IR reconciliation report, period close UI, audit log surfaces, "
             "outstanding payables view.", "Not started"],
        ])

    # Key metrics
    add_h(doc, "Build Metrics", 1)
    add_p(doc, "Engineering output to date on the finance module:")
    add_bullet(doc, "11 backend Alembic migrations applied (finance MySQL schema).")
    add_bullet(doc, "13 MySQL tables managed by the finance service.")
    add_bullet(doc, "21 finance API endpoints exposed.")
    add_bullet(doc, "8 finance UI pages live in the user portal.")
    add_bullet(doc, "227 GL accounts seeded as the default UAE agri-business CoA.")
    add_bullet(doc, "5 default approval rules seeded for the demo company.")
    add_bullet(doc, "Multiple verified end-to-end flows: vendor sync, item sync, PR submit→approve, PO submit→approve→send, GR create→post→JE.")

    # Deferred
    add_h(doc, "Deliberately Deferred to Version 2", 1)
    add_p(doc, "These features were considered and deferred to keep the v1 cycle focused on the "
              "single-currency UAE procure-to-pay flow. Each is scoped but not built.")
    add_bullet(doc, "Multi-currency operations (FX gain/loss accounts).")
    add_bullet(doc, "Vendor returns and credit memos.")
    add_bullet(doc, "Advance payments and prepayments.")
    add_bullet(doc, "Recurring invoices for utilities, rent, etc.")
    add_bullet(doc, "Cash purchases that bypass the AP cycle.")
    add_bullet(doc, "Multi-warehouse separate inventory accounts.")
    add_bullet(doc, "Withholding tax on services.")
    add_bullet(doc, "Customer-facing sales / AR cycle (defer to v2).")
    add_bullet(doc, "Budget tracking and variance reporting.")
    add_bullet(doc, "Period-close enforcement and audit lock.")

    # Risks
    add_h(doc, "Risks and Open Items", 1)
    add_table(doc,
        headers=["Item", "Note"],
        rows=[
            ["Production deployment of finance stack",
             "Currently only running locally. A production deployment plan (MySQL "
             "host, secrets, backup, monitoring) is needed before customer rollout."],
            ["Three-way match design",
             "The design for Phase C (AP Invoice) needs sign-off on tolerance "
             "thresholds (price variance, quantity variance) before build."],
            ["Per-item GL mapping coverage",
             "Each newly created purchase item requires manual finance-side "
             "configuration. Bulk-assign tooling can be added if onboarding "
             "becomes a bottleneck."],
            ["Posting setup completeness",
             "Five default GL accounts must be selected before posting can run. "
             "Implementation guards against running with incomplete setup but the "
             "controller must complete this onboarding step on day one."],
        ])

    # Next steps
    add_h(doc, "Recommended Next Steps", 1)
    add_p(doc, "Sequence for the remaining build, in priority order:")
    add_bullet(doc,
        "Phase C — AP Invoice module with three-way match. This unlocks the "
        "second journal entry of the cycle (recognising the specific vendor "
        "liability and reclaimable VAT).")
    add_bullet(doc,
        "Phase D — Payment module. Closes the cycle by debiting accounts payable "
        "and crediting the bank account.")
    add_bullet(doc,
        "Phase E — operability polish: GR/IR reconciliation report (essential "
        "for month-end), period-close UI, and the outstanding-payables list.")
    add_bullet(doc,
        "Production deployment plan: scope MySQL hosting, backups, secrets "
        "management, and monitoring before customer rollout.")
    add_bullet(doc,
        "Customer onboarding playbook: documented walkthrough for posting "
        "setup, CoA review, and item mapping for new tenants.")

    return doc


# ──────────────────────────────────────────────────────────────────────────────
# 2. USER GUIDE
# ──────────────────────────────────────────────────────────────────────────────

def build_user_guide() -> Document:
    doc = Document()
    style_doc(doc)

    # Title
    title = doc.add_heading("A64 Core Platform — Finance Module User Guide", level=0)
    for run in title.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x36, 0x56)
    sub = doc.add_paragraph()
    sub_run = sub.add_run("How to use the finance module from day one")
    sub_run.italic = True
    meta = doc.add_paragraph()
    meta.add_run("Version: v1 (Phase B)    ·    Date: 2026-05-20").italic = True

    # Welcome
    add_h(doc, "Welcome", 1)
    add_p(doc,
        "This guide walks you through the finance module step by step, from "
        "initial configuration to recording your first goods receipt and "
        "seeing the matching journal entry post automatically. No accounting "
        "background is assumed.")

    # Before you begin
    add_h(doc, "Before You Begin", 1)
    add_p(doc, "You will need:")
    add_bullet(doc, "An account with one of these roles: finance_admin, admin, super_admin "
                    "(controller-level access). Some pages are visible to additional roles "
                    "such as accountant and auditor in read-only mode.")
    add_bullet(doc, "Access to the user portal at the URL provided by your administrator.")
    add_bullet(doc, "Authentication credentials. Default super-admin in the development "
                    "environment is admin@a64platform.com.")

    # Navigation
    add_h(doc, "Finding the Finance Pages", 1)
    add_p(doc, "After logging in, look at the left sidebar. The Finance section appears "
              "as an expandable group with the ledger icon (📒). Click it to expand. "
              "You will see the following pages, in this order:")
    add_table(doc,
        headers=["Page", "When you use it"],
        rows=[
            ["Chart of Accounts",
             "View, search, and manage the list of GL accounts that make up "
             "your company's books. Set up on day one; consulted regularly."],
            ["Approval Rules",
             "Decide which documents (purchase requests, purchase orders, "
             "invoices, payments) require manager or controller approval and "
             "at which amount thresholds. Set up on day one."],
            ["Posting Setup",
             "Tell the system which GL account to use as the default for each "
             "kind of accounting event (vendor payables, bank, VAT, retained "
             "earnings, etc.). Mandatory before any auto-posting can run."],
            ["Item GL Mapping",
             "Assign each purchase item to its inventory and cost-of-sales GL "
             "accounts. Per-item; auto-suggested where possible."],
            ["Journal Entries",
             "Read-only register of every posted journal entry. Shows debits, "
             "credits, and links back to the originating operational document."],
            ["P&L Statement",
             "Profit-and-loss snapshot. Reports module."],
            ["Incoming Preview",
             "See operational documents that are still in Pending Approval, so "
             "finance has visibility into what is about to flow through."],
        ])

    # Day-one setup
    add_h(doc, "Day-One Setup Checklist", 1)
    add_p(doc, "Complete these steps in order. Skipping a step will block downstream "
              "transactions.")

    add_h(doc, "1. Review the Chart of Accounts", 2)
    add_p(doc, "Navigate to Finance → Chart of Accounts. The system seeds 227 accounts "
              "appropriate for a UAE agricultural business. Review them with your "
              "accountant. Add any missing accounts using the New Account button.")
    add_p(doc, "Important fields when creating an account:")
    add_bullet(doc, "Account Number — the unique identifier. Stays with the account "
                    "for its lifetime once postings have happened.")
    add_bullet(doc, "Drawer — one of the nine top-level groups (Assets, Liabilities, "
                    "Equity, Revenue, Cost of Sales, Operating Cost, Non-Operating, "
                    "Other Income, Taxation).")
    add_bullet(doc, "Account Type — broad classification (asset, liability, equity, "
                    "revenue, expense).")
    add_bullet(doc, "Account Level — 'active' for accounts you can post to, 'title' "
                    "and 'drawer' for header-only rows.")
    add_bullet(doc, "Parent Account — establishes hierarchy.")

    add_h(doc, "2. Configure Approval Rules", 2)
    add_p(doc, "Navigate to Finance → Approval Rules. Four default rules come "
              "pre-configured for company 1000. Review them and adjust thresholds.")
    add_p(doc, "Each rule has these fields:")
    add_bullet(doc, "Doc Type — PR, PO, AP_INVOICE, OUTGOING_PAYMENT, or others.")
    add_bullet(doc, "Company Code — usually 1000 in single-company setups.")
    add_bullet(doc, "Approver Role — who must approve (e.g. procurement_manager).")
    add_bullet(doc, "Approval Mode — either 'Always required' or 'Above threshold'.")
    add_bullet(doc, "Threshold Amount — only meaningful in 'Above threshold' mode.")
    add_bullet(doc, "Priority — lower number = higher priority when multiple rules match.")
    add_p(doc, "Use the Test Resolution widget at the bottom to verify a given document "
              "type + amount combination triggers the rule you expect.")

    add_h(doc, "3. Complete the Posting Setup", 2)
    add_p(doc, "Navigate to Finance → Posting Setup. This is the most important "
              "configuration step — until it is complete, the posting engine cannot "
              "create journal entries.")
    add_p(doc, "Required fields (marked with *):")
    add_table(doc,
        headers=["Field", "Recommended choice (seeded CoA)"],
        rows=[
            ["AP Control Account *", "221000-001 Trade Payables - Suppliers"],
            ["Bank Account *", "126000-002 Cash at Bank - AED Operating"],
            ["GR/IR Clearing Account *", "221000-002 Goods Received Not Invoiced"],
            ["Input VAT Account *", "122000-001 Input VAT Recoverable"],
            ["Retained Earnings Account *", "312000-001 Retained Earnings - Prior Years"],
        ])
    add_p(doc, "Optional fields (set when ready):")
    add_table(doc,
        headers=["Field", "Recommended choice"],
        rows=[
            ["Cash Account", "126000-001 Petty Cash"],
            ["Output VAT Account", "222000-001 Output VAT Payable"],
            ["AR Control Account", "Leave blank in v1 (sales side not yet built)"],
            ["Purchase Price Variance Account",
             "Leave blank in v1 (add when AP Invoice ships in Phase C)"],
            ["Rounding Account", "Leave blank in v1"],
        ])
    add_p(doc, "Click Save Configuration. The completeness badge at the top should turn "
              "green ('Configured') once all five required fields are set.")

    add_h(doc, "4. Map Items to GL Accounts", 2)
    add_p(doc, "Navigate to Finance → Item GL Mapping. This page shows every purchase "
              "item your operational team has created. For each item assign the GL "
              "accounts the system will use when goods are received and consumed.")
    add_p(doc, "Inventory Account is the asset account the item lands in when received. "
              "For raw materials, the system auto-suggests "
              "'121000-002 Raw Materials - Fertilisers' but you should override per item — "
              "for example, choose 'Raw Materials - Seeds' for a seed item.")
    add_p(doc, "COGS Account is the expense account used when the item is consumed or "
              "sold. Typically a 511000-series Cost-of-Sales account corresponding to "
              "the material category.")
    add_p(doc, "Valuation Method defaults to Moving Average — keep this unless your "
              "accountant has a reason to choose Standard Cost or FIFO.")
    add_p(doc, "Save each row, or use the Save All button to commit multiple rows at once.")

    # Daily flow
    add_h(doc, "The Procure-to-Pay Cycle in Practice", 1)
    add_p(doc, "Once configured, the day-to-day flow that produces journal entries is "
              "as follows. Operational staff handle the first three steps; the system "
              "produces the journal entries automatically.")

    add_h(doc, "Step 1: Purchase Request (operational team)", 2)
    add_p(doc, "Operations → Purchasing → Purchase Requests → New PR.")
    add_p(doc, "A user raises an internal request listing items and quantities needed. "
              "Submit for approval. No accounting entry is created at this stage.")

    add_h(doc, "Step 2: Purchase Order (operational team)", 2)
    add_p(doc, "From the approved PR, click 'Create PO from PR' or go to "
              "Operations → Purchasing → Purchase Orders → New PO. Pick the vendor, "
              "confirm prices, submit. After approval, click Send to commit the order "
              "to the vendor. Still no accounting entry — a PO is a commitment, not "
              "yet a transaction.")

    add_h(doc, "Step 3: Goods Receipt (warehouse / receiving team)", 2)
    add_p(doc, "Operations → Goods Receipts → New from PO.")
    add_p(doc, "When goods physically arrive, the receiving team records what was "
              "actually received. Each line defaults to the remaining quantity from "
              "the PO but can be reduced for partial deliveries.")
    add_p(doc, "Click Post on the Draft GR. This is the magic moment:")
    add_bullet(doc, "The GR transitions to Posted (no further edits).")
    add_bullet(doc, "The parent PO's open quantity decrements. If fully received, "
                    "the PO auto-closes.")
    add_bullet(doc, "An event is sent to the finance service.")
    add_bullet(doc, "Finance creates a journal entry within a few seconds: "
                    "debit the inventory account configured on the item, credit the "
                    "GR/IR clearing account configured in Posting Setup.")

    add_h(doc, "Step 4: View the Journal Entry", 2)
    add_p(doc, "After posting a GR you will see a green banner with a link 'View "
              "Journal Entry →'. Click it to navigate to Finance → Journal Entries "
              "with the relevant GR number pre-filtered.")
    add_p(doc, "Click the row to expand it. You will see two or more lines: the "
              "debits to your inventory accounts and the matching credit to "
              "GR/IR Clearing. The total debits equal total credits, as they must.")

    # Sample journal entries
    add_h(doc, "What the Journal Entries Look Like", 1)
    add_p(doc, "Concrete example: ten bags of fertiliser at AED 100 each.")
    add_h(doc, "At Goods Receipt (built today, Phase B)", 2)
    add_code(doc, "DR  121000-002 Raw Materials - Fertilisers   1,000.00\n"
                  "CR  221000-002 Goods Received Not Invoiced              1,000.00")
    add_p(doc, "VAT is not recognised at this stage. The system only knows the "
              "physical quantity received and the PO price.")

    add_h(doc, "At Vendor Invoice (coming in Phase C)", 2)
    add_code(doc, "DR  221000-002 Goods Received Not Invoiced     1,000.00\n"
                  "DR  122000-001 Input VAT Recoverable               50.00\n"
                  "CR  221000-001 Trade Payables - Suppliers                1,050.00")
    add_p(doc, "The temporary GR/IR holding is cleared. The reclaimable VAT becomes "
              "an asset. The specific vendor liability is recognised.")

    add_h(doc, "At Payment (coming in Phase D)", 2)
    add_code(doc, "DR  221000-001 Trade Payables - Suppliers       1,050.00\n"
                  "CR  126000-002 Cash at Bank - AED Operating                 1,050.00")
    add_p(doc, "Money leaves the bank. Vendor liability is cleared. Cycle closed.")

    # Common scenarios
    add_h(doc, "Common Scenarios and Tips", 1)

    add_h(doc, "Partial Deliveries", 2)
    add_p(doc, "If you receive 8 of 10 bags, post a GR for 8. The PO line remains "
              "open for the remaining 2 bags. A second GR can be created later for "
              "the balance, which posts a second journal entry.")

    add_h(doc, "Items With No GL Mapping", 2)
    add_p(doc, "On Item GL Mapping, items without an Inventory Account that are "
              "expected to need one (raw materials, consumables) are highlighted in "
              "amber. These items will block GR posting until you assign an account.")

    add_h(doc, "Wrong Inventory Account Assigned", 2)
    add_p(doc, "Go to Finance → Item GL Mapping, find the item, change the "
              "Inventory Account, save. New goods receipts will use the new account. "
              "Previously posted journal entries are not changed (immutable by design).")

    add_h(doc, "Posted GR Was Wrong", 2)
    add_p(doc, "Posted goods receipts cannot be edited or deleted in v1. If a "
              "mistake is significant, contact your finance lead — the v1 workaround "
              "is a manual journal adjustment. Phase B+ will add a 'reverse GR' "
              "action for proper corrections.")

    add_h(doc, "Approval Not Triggering", 2)
    add_p(doc, "Use the Test Resolution widget on Finance → Approval Rules. Enter "
              "the document type and amount; the widget shows you whether approval "
              "will be required and which role would receive it.")

    # Glossary
    add_h(doc, "Glossary", 1)
    add_table(doc,
        headers=["Term", "Meaning"],
        rows=[
            ["GL / General Ledger", "The master record of every accounting entry. "
             "Made up of accounts grouped into drawers (Assets, Liabilities, etc.)."],
            ["Chart of Accounts (CoA)", "The list of accounts that make up your GL. "
             "Each account has a number, name, and classification."],
            ["Journal Entry (JE)", "One balanced accounting transaction made of "
             "two or more lines. Total debits always equal total credits."],
            ["Debit (DR) / Credit (CR)", "The two sides of every accounting entry. "
             "They mean different things depending on the account type — for assets "
             "and expenses, debit increases; for liabilities, equity, and revenue, "
             "credit increases."],
            ["PR / PO / GR", "Purchase Request / Purchase Order / Goods Receipt — the "
             "three operational document stages before the vendor's invoice arrives."],
            ["GR/IR Clearing", "Goods Received / Invoice Received. A temporary "
             "liability account used between the moment goods arrive and the moment "
             "the vendor's invoice is recorded."],
            ["AP / Accounts Payable", "What we owe vendors. A liability."],
            ["VAT", "Value-Added Tax. In the UAE, the standard rate is 5%. Input "
             "VAT (on purchases) is reclaimable; output VAT (on sales) is payable."],
            ["Fiscal Period", "An accounting time-window, usually a month. Postings "
             "can only happen against an Open period."],
            ["Control Account", "A GL account managed by the system (such as Trade "
             "Payables). Not directly editable by users to maintain integrity."],
            ["Three-way Match", "The check at vendor-invoice time that the invoice "
             "agrees with the original PO and the actual GR (quantity and price). "
             "Comes in Phase C."],
        ])

    # Closing
    add_h(doc, "Support", 1)
    add_p(doc, "If a page reports an error or behaves unexpectedly, capture the "
              "page URL, the error text, and the time. Forward to your platform "
              "administrator with that information.")
    add_p(doc, "Additional reference documents are stored under "
              "Docs/4-Finance-Mod-docs/ in the project repository.")

    return doc


# ──────────────────────────────────────────────────────────────────────────────
# Main
# ──────────────────────────────────────────────────────────────────────────────

def main() -> None:
    pm_doc = build_pm_report()
    pm_path = OUT_DIR / "FINANCE_PM_REPORT.docx"
    pm_doc.save(str(pm_path))
    print(f"wrote {pm_path}")

    user_doc = build_user_guide()
    user_path = OUT_DIR / "FINANCE_USER_GUIDE.docx"
    user_doc.save(str(user_path))
    print(f"wrote {user_path}")


if __name__ == "__main__":
    main()
