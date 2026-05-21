"""
Generate FINANCE_PM_REPORT_POST_D.docx — a delta PM report covering work
done in the session that followed the Phase C+D report.

Scope of this delta:
  - Reversing-entry pattern migration (and the two bugs it surfaced).
  - Trial Balance UX simplification.
  - New Tenant Bootstrap Wizard (operational feature for fresh deployments).
  - Data-integrity post-mortems and cleanup.
  - What still remains for v1.

The file lands in Docs/4-Finance-Mod-docs/ alongside FINANCE_PM_REPORT_PHASE_C_D.docx.
"""

from pathlib import Path

from docx import Document
from docx.shared import Cm, Pt, RGBColor

OUT = Path(__file__).resolve().parent.parent / "Docs" / "4-Finance-Mod-docs"
OUT.mkdir(parents=True, exist_ok=True)


def style_doc(doc):
    n = doc.styles["Normal"]
    n.font.name = "Calibri"
    n.font.size = Pt(11)
    n.paragraph_format.space_before = Pt(0)
    n.paragraph_format.space_after = Pt(6)
    n.paragraph_format.line_spacing = 1.25


def add_h(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.color.rgb = RGBColor(0x1F, 0x36, 0x56)


def add_p(doc, text, italic=False, bold=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.italic = italic
    r.bold = bold


def add_bullet(doc, text):
    p = doc.add_paragraph(text, style="List Bullet")
    p.paragraph_format.left_indent = Cm(0.75)


def add_code(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = "Consolas"
    r.font.size = Pt(10)
    p.paragraph_format.left_indent = Cm(0.5)


def add_table(doc, headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Light Grid Accent 1"
    for i, h in enumerate(headers):
        t.rows[0].cells[i].text = h
        for run in t.rows[0].cells[i].paragraphs[0].runs:
            run.bold = True
    for ri, row in enumerate(rows, start=1):
        for ci, val in enumerate(row):
            t.rows[ri].cells[ci].text = val
            for p in t.rows[ri].cells[ci].paragraphs:
                p.paragraph_format.space_after = Pt(2)
    doc.add_paragraph()


# ============================================================================
# Report body
# ============================================================================

def build_report():
    doc = Document()
    style_doc(doc)

    title = doc.add_heading(
        "Finance Module — Post-Phase D Delta Report",
        level=0,
    )
    for r in title.runs:
        r.font.color.rgb = RGBColor(0x1F, 0x36, 0x56)

    sub = doc.add_paragraph()
    sub.add_run(
        "Reversing-entry migration, Trial Balance cleanup, Tenant Bootstrap "
        "Wizard, and data-integrity work"
    ).italic = True
    meta = doc.add_paragraph()
    meta.add_run(
        "Author: Viet Anh   ·   Period: post-Phase D session   ·   "
        "Date: 2026-05-21"
    ).italic = True
    doc.add_paragraph()

    # ─── Executive summary ─────────────────────────────────────────────────
    add_h(doc, "Executive Summary", 1)
    add_p(
        doc,
        "This report continues from FINANCE_PM_REPORT_PHASE_C_D.docx. It "
        "covers the work completed in the session immediately following "
        "the Phase C+D delivery. Four distinct streams of work landed:",
    )
    add_bullet(
        doc,
        "Reversing-entry pattern migration — the JE reverse action was "
        "rebuilt to follow the GAAP/IFRS standard used by SAP, NetSuite, "
        "and QuickBooks. Two underlying bugs were surfaced and fixed in "
        "the process (a sub-ledger orphan and a double-cancellation in "
        "the reports).",
    )
    add_bullet(
        doc,
        "Trial Balance UX simplification — the redundant Balance column "
        "was removed; the page now reads cleanly as a pure debit/credit "
        "check. Behavioural note on gross totals under reversing entries "
        "documented for accountants.",
    )
    add_bullet(
        doc,
        "Tenant Bootstrap Wizard — a new 6-step UI flow that walks a "
        "super_admin through organisation creation, self-assignment, "
        "finance company creation (which auto-seeds 230 GL accounts + "
        "5 tax codes + approval rules), and first fiscal period. Solves "
        "the \"deployed to a fresh server, finance won't load\" problem.",
    )
    add_bullet(
        doc,
        "Data-integrity work — two real production-grade issues were "
        "diagnosed via live data: an orphaned credit on the Vendor "
        "Sub-Ledger, and a phantom Trial Balance figure after a JE "
        "reversal. Both root-caused and fixed; existing wrong rows "
        "corrected in place via transactional DB patches.",
    )
    add_p(
        doc,
        "The cumulative effect is that the finance module is now safe to "
        "ship to a clean server: a new tenant can be brought to a "
        "working state through the UI alone, with no SQL or seed scripts. "
        "The reversing-entry pattern brings the accounting behaviour to "
        "parity with mid-market ERPs.",
    )

    # ─── Stream 1: Reversing-Entry Migration ───────────────────────────────
    add_h(doc, "Stream 1 — Reversing-Entry Pattern Migration", 1)

    add_h(doc, "What was wrong", 2)
    add_p(
        doc,
        "When a user reversed a vendor payment, the implementation was "
        "doing two opposing things at once: it set the original journal "
        "entry to status='void' AND posted a separate reversal JE that "
        "swapped debit and credit. Reports filter out voided JEs by "
        "default, so the original was excluded while the reversal was "
        "counted — producing a net effect that double-cancelled the "
        "transaction.",
    )
    add_p(
        doc,
        "A second, subtler bug compounded the problem: the reversal "
        "lines were overwriting referenceLineId (the sub-ledger key — "
        "vendorId on AP lines) with the original line's primary key, "
        "intended for traceability. The Vendor Sub-Ledger groups on "
        "referenceLineId, so the reversal's credit landed under a UUID "
        "that no vendor owned — appearing as a phantom 35,000 AED "
        "outstanding balance against \"nobody\".",
    )

    add_h(doc, "What was changed", 2)
    add_bullet(
        doc,
        "JE reverse now follows the standard reversing-entry pattern: "
        "the original stays posted; the reversal is its own posted JE; "
        "both live on the books and mathematically net to zero. \"Void\" "
        "is reserved for genuine posting errors (never created by the "
        "reverse action).",
    )
    add_bullet(
        doc,
        "Reversal lines now preserve the original referenceLineId, so "
        "the Vendor Sub-Ledger and any other sub-ledger correctly nets "
        "the entry against the same vendor.",
    )
    add_bullet(
        doc,
        "A guard was added: an attempt to reverse a JE that already has "
        "a reversal JE returns 400 with a clear message naming the "
        "existing reversal.",
    )
    add_bullet(
        doc,
        "API responses (Journal Entries list/detail, Vendor Payments "
        "list/detail) now carry a reversedByJeNumber field, populated "
        "by a single batched lookup per request. Frontend uses this to "
        "decide whether to render the Reversed badge — no extra round "
        "trips per row.",
    )
    add_bullet(
        doc,
        "Frontend UI: the Vendor Payments list shows a red Reversed pill "
        "with a struck-through amount; the Payment Detail page shows a "
        "full banner naming the reversal JE; the Journal Entries list "
        "shows a Reversed status badge with the reverser's JE# in the "
        "tooltip and disables the Reverse Entry button on already-"
        "reversed rows.",
    )
    add_bullet(
        doc,
        "The reverse mutation invalidates downstream caches (payments, "
        "ap-totals-paid, vendor sub-ledger, AP aging, trial balance) "
        "so all balances refresh without a manual reload.",
    )

    add_h(doc, "Data migration", 2)
    add_p(
        doc,
        "Two transactional patches ran against the live finance database:",
    )
    add_bullet(
        doc,
        "Un-voided every JE that had been voided by a prior reverse "
        "action while also having a corresponding reversal JE — one "
        "row matched in this environment.",
    )
    add_bullet(
        doc,
        "Patched the referenceLineId on the orphan reversal lines to "
        "match the original lines' vendorId — two lines corrected.",
    )
    add_p(
        doc,
        "After migration, the Vendor Sub-Ledger total outstanding "
        "balance dropped from 35,000 to 0 for the affected vendor, "
        "matching the real-world position.",
    )

    add_h(doc, "Why this matters", 2)
    add_p(
        doc,
        "The reversing-entry pattern is what SAP, NetSuite, Oracle, and "
        "QuickBooks all implement. An auditor expecting GAAP/IFRS-"
        "compliant behaviour will see a complete history: original "
        "entry, reversal entry, both visible, both posted, net zero. "
        "Voids should be rare and require an explicit reason. The "
        "previous behaviour would have been a finding in any audit.",
    )

    # ─── Stream 2: Trial Balance Cleanup ───────────────────────────────────
    add_h(doc, "Stream 2 — Trial Balance UX Simplification", 1)

    add_h(doc, "Change", 2)
    add_bullet(
        doc,
        "Removed the Balance column from the per-account table and the "
        "totals footer. The page now shows Account Number, Account "
        "Name, Drawer, Total Debit, Total Credit — the canonical Trial "
        "Balance layout.",
    )
    add_bullet(
        doc,
        "Drawer-section colSpan adjusted; footer placeholder cell "
        "removed; export contract unchanged (Balance was never in any "
        "export).",
    )

    add_h(doc, "Behavioural note on gross totals", 2)
    add_p(
        doc,
        "Under the new reversing-entry pattern, the Trial Balance's "
        "GROSS debit and credit totals will grow by 2× the reversed "
        "amount each time a JE is reversed (the original's debit + the "
        "reversal's matching debit, same for credit). NET per-account "
        "balances remain correct — the two entries cancel. This is the "
        "same behaviour as in SAP, NetSuite, and QuickBooks, and is the "
        "expected GAAP/IFRS view. The accountant should not be alarmed "
        "to see totals tick up after a reversal; they should verify "
        "that net per-account balances on Trial Balance still tie to "
        "Vendor Sub-Ledger / AP Aging.",
    )

    # ─── Stream 3: Tenant Bootstrap Wizard ─────────────────────────────────
    add_h(doc, "Stream 3 — Tenant Bootstrap Wizard (new operational feature)", 1)

    add_h(doc, "The problem", 2)
    add_p(
        doc,
        "When deploying to a clean server, the bootstrap super_admin "
        "account is created without an organizationId. Every finance "
        "page is gated on a non-null organizationId, so the entire "
        "module silently fails to load. There was no UI path to create "
        "an organisation, assign the user, or seed the per-org finance "
        "master data — only direct API calls or SQL.",
    )

    add_h(doc, "The solution", 2)
    add_p(
        doc,
        "A 6-step wizard at /admin/tenant-setup, accessible without an "
        "organizationId, walks the super_admin through:",
    )
    add_table(
        doc,
        ["Step", "What happens", "Backend"],
        [
            ["0  Welcome", "Detects current state, decides which steps to skip.",
             "GET /organizations, GET /finance/companies"],
            ["1  Organisation", "Pick existing or create new (name, slug, industries, logoUrl).",
             "POST /organizations/ (super_admin)"],
            ["2  Self-assign", "Assigns the chosen org to the current user; session refreshes silently.",
             "PATCH /admin/users/{id}/organization (new endpoint, super_admin)"],
            ["3  Company code", "Creates the finance company; auto-seeds 230 GL accounts + 5 tax codes + approval rules.",
             "POST /finance/companies"],
            ["4  Fiscal period", "Creates the first open period (defaults to current month). Picker if multiple companies exist.",
             "POST /finance/periods"],
            ["5  Done", "Summary of what was created; CTA to dashboard.", "—"],
        ],
    )

    add_h(doc, "Design properties", 2)
    add_bullet(
        doc,
        "Auto-skip: a returning super_admin who already has an org "
        "lands on the first incomplete step (e.g. company code missing "
        "→ step 3). Completed steps are shown with green checks.",
    )
    add_bullet(
        doc,
        "Auto-redirect: an org-less super_admin who visits any other "
        "page is redirected to the wizard. Non-super-admins without an "
        "org get a clear message asking them to contact a super_admin "
        "— they cannot self-fix.",
    )
    add_bullet(
        doc,
        "Session refresh without logout: after assignment, the auth "
        "store calls GET /auth/me and silently updates the user object. "
        "No re-login required.",
    )
    add_bullet(
        doc,
        "Permission tightening: POST /organizations/ was previously "
        "open to admin OR super_admin; it is now super_admin only. "
        "Regular admins operate inside an org but cannot create new "
        "tenants. The new PATCH /admin/users/{id}/organization is also "
        "super_admin only.",
    )
    add_bullet(
        doc,
        "Smart defaults on the company step: companyCode \"1000\", "
        "currency AED, valuation MovingAverage. Pre-filled but editable.",
    )
    add_bullet(
        doc,
        "Multi-company picker on the fiscal-period step when multiple "
        "companies exist; default selection prefers the just-created "
        "code, then \"1000\", then first sorted — so the user is never "
        "silently dropped into an unexpected company.",
    )

    add_h(doc, "Tenant Bootstrap — verified outcomes", 2)
    add_bullet(doc, "Org creation returns 201 with the new organisation row.")
    add_bullet(doc, "User assignment returns the updated user including organizationId.")
    add_bullet(doc, "GET /auth/me reflects the new orgId on the next request — no re-login.")
    add_bullet(
        doc,
        "POST /finance/companies on a fresh org seeds 230 GL accounts "
        "and 5 tax codes, confirmed via direct DB queries.",
    )
    add_bullet(doc, "POST /finance/periods creates the first period in status=Open.")
    add_bullet(doc, "Chart of Accounts page loads correctly with the seeded accounts.")

    # ─── Stream 4: Data-Integrity Post-Mortems ─────────────────────────────
    add_h(doc, "Stream 4 — Data-Integrity Post-Mortems", 1)

    add_h(doc, "Phantom Vendor Sub-Ledger balance", 2)
    add_p(
        doc,
        "Symptom: Vendor Sub-Ledger showed a 35,000 AED balance owed "
        "to a vendor whose UUID matched no entity. Root cause: the JE "
        "reversal code was overwriting referenceLineId (the vendor "
        "sub-ledger key) with the original line's jeLineId. Fixed in "
        "code and patched in the live DB. Documented in Stream 1.",
    )

    add_h(doc, "Trial Balance jumped after reversal", 2)
    add_p(
        doc,
        "Symptom: After data migration to the reversing-entry pattern, "
        "the Trial Balance gross totals moved from 140,000 to 175,000 "
        "while net balances stayed correct. Root cause: not a bug — "
        "this is the expected effect of un-voiding the original JE "
        "(so it is once again counted in gross totals) while keeping "
        "the reversal posted. Explained to the user and documented in "
        "Stream 2's behavioural note for future accountants.",
    )

    add_h(doc, "Dev-environment data pollution", 2)
    add_p(
        doc,
        "During end-to-end verification of the Tenant Bootstrap Wizard, "
        "an automated test agent ran the wizard against the live "
        "development tenant. The user's super_admin account was "
        "reassigned to a test organisation, hiding their real data "
        "from the finance UI. Resolved via transactional cleanup:",
    )
    add_bullet(doc, "User org reassignment reverted in Mongo (one super_admin restored).")
    add_bullet(doc, "Two test organisations deleted from Mongo.")
    add_bullet(
        doc,
        "Two test company codes deleted from MySQL along with their "
        "seeded data (460 GL accounts, 10 tax codes, 8 approval rules, "
        "2 fiscal periods).",
    )
    add_p(
        doc,
        "Process note: future automated UI tests should run against an "
        "isolated worktree + ephemeral DB, or limit themselves to read-"
        "only verification with the live test reserved for the user. "
        "Adjusting agent invocation patterns accordingly.",
    )

    # ─── What's owed ────────────────────────────────────────────────────────
    add_h(doc, "What still remains for v1", 1)

    add_p(doc, "Deferred from Phase E and unchanged by this delta:", italic=True)
    add_bullet(doc, "Accounts Receivable (sales-side) — the major remaining functional area.")
    add_bullet(
        doc,
        "MFA/TOTP and SSO/OAuth — login factors still pending; super_admin "
        "is currently single-factor.",
    )
    add_bullet(
        doc,
        "Field-level encryption for sensitive PII (emiratesId, salary) — "
        "currently stored in cleartext at rest.",
    )
    add_bullet(doc, "WAF, MongoDB replica set, SIEM, penetration testing.")
    add_bullet(
        doc,
        "Item GL Mapping bulk-edit and richer error messages — known "
        "small ergonomics items.",
    )
    add_bullet(
        doc,
        "P&L Statement and additional management reports beyond Trial "
        "Balance / Vendor Sub-Ledger / AP Aging.",
    )

    add_h(doc, "New items surfaced by this session", 2)
    add_bullet(
        doc,
        "Division-aware org assignment: the new PATCH /admin/users/"
        "{id}/organization accepts optional divisionAccess and "
        "defaultDivisionId, but the wizard does not yet surface them. "
        "A division-management UI is a natural next step once the "
        "first tenant has more than one operational unit.",
    )
    add_bullet(
        doc,
        "Bulk org assignment: a super_admin can only assign users one "
        "at a time. A bulk-assign affordance in User Management would "
        "speed up onboarding for larger tenants.",
    )
    add_bullet(
        doc,
        "Reverse-action audit log: the reversal JE's description "
        "carries the reason text, but no structured audit table records "
        "who reversed what and when in a queryable form. The audit_log "
        "table exists but the reverse endpoint does not yet write to it.",
    )

    # ─── Closing ───────────────────────────────────────────────────────────
    add_h(doc, "Conclusion", 1)
    add_p(
        doc,
        "The finance module is now (a) accounting-correct under the "
        "industry-standard reversing-entry pattern, (b) deployable to "
        "a clean server through the UI alone, and (c) safe against the "
        "two specific data-integrity bugs that surfaced this session. "
        "The Phase A → D.5 functional scope is unchanged; what landed "
        "in this delta is robustness, correctness, and operational "
        "readiness for new tenants.",
    )

    add_h(doc, "Reference Documents", 1)
    add_bullet(doc, "FINANCE_PM_REPORT_PHASE_C_D.docx — preceding report.")
    add_bullet(doc, "FINANCE_USER_GUIDE.docx — end-user walkthrough (Phase A → D.5).")
    add_bullet(doc, "POSTING_ENGINE_ROADMAP.md — phased build plan + controls-debt register.")
    add_bullet(doc, "INTEGRATION_MODEL.md — cross-service event design and immutability rules.")
    add_bullet(doc, "FEEDBACK_RESPONSE.docx — itemised response to the prior PM feedback note.")

    return doc


def main():
    doc = build_report()
    out_path = OUT / "FINANCE_PM_REPORT_POST_D.docx"
    doc.save(str(out_path))
    print(f"wrote {out_path}")


if __name__ == "__main__":
    main()
