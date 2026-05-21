"""
Generate two updated finance-module .docx deliverables for Phase C + D:

  1. FINANCE_USER_GUIDE.docx — OVERWRITES the existing user guide with the
     prior Phase A/B content PLUS new Phase C (AP Invoice + reverse-charge
     VAT) and Phase D (Payment) sections, plus Phase D.5 (Period close).

  2. FINANCE_PM_REPORT_PHASE_C_D.docx — NEW delta-focused report covering
     what was built in Phases C and D, the PM-feedback items addressed,
     and what remains.

Both files land in Docs/4-Finance-Mod-docs/.
"""

from pathlib import Path

from docx import Document
from docx.shared import Cm, Pt, RGBColor

OUT = Path(__file__).resolve().parent.parent / "Docs" / "4-Finance-Mod-docs"
OUT.mkdir(parents=True, exist_ok=True)


def style_doc(doc):
    n = doc.styles["Normal"]
    n.font.name = "Calibri"
    n.font.size = Pt(11)
    n.paragraph_format.space_before = Pt(0)
    n.paragraph_format.space_after = Pt(6)
    n.paragraph_format.line_spacing = 1.25


def add_h(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.color.rgb = RGBColor(0x1F, 0x36, 0x56)


def add_p(doc, text, italic=False, bold=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.italic = italic
    r.bold = bold


def add_bullet(doc, text):
    p = doc.add_paragraph(text, style="List Bullet")
    p.paragraph_format.left_indent = Cm(0.75)


def add_code(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = "Consolas"
    r.font.size = Pt(10)
    p.paragraph_format.left_indent = Cm(0.5)


def add_table(doc, headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Light Grid Accent 1"
    for i, h in enumerate(headers):
        t.rows[0].cells[i].text = h
        for run in t.rows[0].cells[i].paragraphs[0].runs:
            run.bold = True
    for ri, row in enumerate(rows, start=1):
        for ci, val in enumerate(row):
            t.rows[ri].cells[ci].text = val
            for p in t.rows[ri].cells[ci].paragraphs:
                p.paragraph_format.space_after = Pt(2)
    doc.add_paragraph()


# ============================================================================
# 1. UPDATED USER GUIDE (full content, Phase A through D inclusive)
# ============================================================================

def build_user_guide():
    doc = Document()
    style_doc(doc)

    t = doc.add_heading("A64 Core Platform — Finance Module User Guide", level=0)
    for r in t.runs:
        r.font.color.rgb = RGBColor(0x1F, 0x36, 0x56)
    sub = doc.add_paragraph()
    sub.add_run("How to use the finance module end-to-end").italic = True
    meta = doc.add_paragraph()
    meta.add_run("Version: v1 (Phases A → D.5)    ·    Updated: 2026-05-21").italic = True

    add_h(doc, "Welcome", 1)
    add_p(doc,
        "This guide walks you through the finance module step by step, from "
        "initial setup to the complete procure-to-pay cycle: raising a purchase "
        "request, receiving goods, recording the vendor's invoice, paying it, "
        "and checking that the books balance. No accounting background is "
        "assumed.")

    add_h(doc, "Before You Begin", 1)
    add_bullet(doc, "An account with one of these roles: finance_admin, admin, "
                    "super_admin (full access). Other roles like accountant, "
                    "auditor, and procurement_* see relevant pages with "
                    "restricted permissions.")
    add_bullet(doc, "Access to the user portal at the URL provided by your administrator.")
    add_bullet(doc, "Authentication credentials. Default super-admin in the "
                    "development environment is admin@a64platform.com.")

    add_h(doc, "Finding the Pages", 1)
    add_p(doc, "Two sidebar groups carry the finance workflow:")
    add_p(doc, "Operations → Purchasing", bold=True)
    add_bullet(doc, "Vendors, Purchase Items, Payment Terms — master data managed by procurement.")
    add_bullet(doc, "Purchase Requests, Purchase Orders, Goods Receipts, AP Invoices — the document flow.")
    add_bullet(doc, "Approval Inbox — your pending approvals.")
    add_p(doc, "Finance", bold=True)
    add_bullet(doc, "Chart of Accounts, Approval Rules, Posting Setup, Item GL Mapping — controller setup.")
    add_bullet(doc, "Journal Entries, Trial Balance, P&L Statement — accounting output.")
    add_bullet(doc, "Vendor Payments, AP Aging, Vendor Sub-Ledger — payables management.")
    add_bullet(doc, "Fiscal Periods, Incoming Preview — period control and operational visibility.")

    # ───────────────────────────────────────────────────────────────────────
    add_h(doc, "Part 1 — Day-One Setup Checklist", 1)
    add_p(doc, "Complete these steps in order. Skipping a step blocks downstream transactions.")

    add_h(doc, "Step 1.1 — Review the Chart of Accounts", 2)
    add_p(doc, "Navigate to Finance → Chart of Accounts. The system seeds 230+ "
              "UAE-aligned accounts. Review with your accountant; add anything "
              "specific to your business via the New Account button.")
    add_p(doc, "Key fields when creating an account:")
    add_bullet(doc, "Account Number — unique identifier, stays with the account.")
    add_bullet(doc, "Drawer — one of the nine top-level groups (Assets, "
                    "Liabilities, Equity, Revenue, Cost of Sales, Operating "
                    "Cost, Non-Operating, Other Income, Taxation).")
    add_bullet(doc, "Account Type — broad classification (asset, liability, equity, revenue, expense).")
    add_bullet(doc, "Account Level — 'active' for postable accounts; 'title' and "
                    "'drawer' for header-only rows.")

    add_h(doc, "Step 1.2 — Configure Approval Rules", 2)
    add_p(doc, "Navigate to Finance → Approval Rules. Default rules come "
              "pre-configured for company 1000. Adjust thresholds and approver "
              "roles for PR, PO, AP_INVOICE, and OUTGOING_PAYMENT documents. "
              "Use the Test Resolution widget at the bottom to verify a "
              "(docType, amount) combination triggers the rule you expect.")

    add_h(doc, "Step 1.3 — Complete the Posting Setup", 2)
    add_p(doc, "Navigate to Finance → Posting Setup. Required before posting "
              "can run. Recommended values from the seeded CoA:")
    add_table(doc,
        headers=["Field", "Recommended account"],
        rows=[
            ["AP Control Account *", "221000-001 Trade Payables - Suppliers"],
            ["Bank Account *", "126000-002 Cash at Bank - AED Operating"],
            ["GR/IR Clearing Account *", "223000-004 Goods Received Not Invoiced"],
            ["Input VAT Account *", "122000-001 Input VAT Recoverable"],
            ["Output VAT Account", "222000-001 Output VAT Payable (required for reverse-charge tax codes)"],
            ["Retained Earnings Account *", "312000-001 Retained Earnings - Prior Years"],
            ["Purchase Price Variance Account", "514000-004 Purchase Price Variance"],
            ["Cash Account", "126000-001 Petty Cash"],
            ["Rounding Account", "617000-011 Rounding Differences"],
            ["Default Valuation Method *", "Moving Average (company-wide, per IAS 2)"],
        ])
    add_p(doc, "Click Save Configuration. The green 'Configured' badge appears "
              "when all required fields are set.")

    add_h(doc, "Step 1.4 — Map Items to GL Accounts", 2)
    add_p(doc, "Navigate to Finance → Item GL Mapping. For each purchase item "
              "the procurement team has created, assign:")
    add_bullet(doc, "Inventory Account — the asset account debited when the "
                    "item is received (e.g. 121000-001 for Seeds).")
    add_bullet(doc, "COGS Account — the cost-of-sales account used when the "
                    "item is consumed (e.g. 511000-001 for Seeds, 511000-002 "
                    "for Fertilisers).")
    add_p(doc, "Use the 'Auto-assign Defaults' button to set all raw-material "
              "items to the system default in one click, then override where "
              "appropriate.")

    add_h(doc, "Step 1.5 — Open the Fiscal Periods", 2)
    add_p(doc, "Navigate to Finance → Fiscal Periods. Click 'Create Periods "
              "for Year...' to bulk-create 12 monthly periods (or 13 for 4-4-5 "
              "calendars, or 4 quarterly). Use 'All Open' for development or "
              "'All Closed except current month' for stricter control.")
    add_p(doc, "Closed periods reject any new journal entries. To correct a "
              "transaction in a closed period, either reopen the period "
              "(audited) or post a reversal in the current open period.")

    # ───────────────────────────────────────────────────────────────────────
    add_h(doc, "Part 2 — The Procure-to-Pay Cycle", 1)
    add_p(doc, "Once configured, day-to-day operation flows through five "
              "stages. Three of them produce real journal entries; the "
              "system creates them automatically.")
    add_table(doc,
        headers=["Stage", "What it is", "Journal entry produced?"],
        rows=[
            ["Purchase Request (PR)", "Internal 'may I buy this' document", "No"],
            ["Purchase Order (PO)", "External order sent to the vendor", "No (accrual accounting)"],
            ["Goods Receipt (GR)", "Warehouse confirms physical delivery", "Yes — JE #1"],
            ["AP Invoice", "Vendor's bill, three-way matched", "Yes — JE #2"],
            ["Payment", "Cash leaves the bank", "Yes — JE #3"],
        ])

    add_h(doc, "Stage A — Purchase Request and Purchase Order", 2)
    add_p(doc, "1. Operations → Purchasing → Purchase Requests → New PR. List items + quantities needed. Submit for approval.")
    add_p(doc, "2. From the approved PR, click 'Create PO from PR' or go to Purchase Orders → New PO. "
              "Pick the vendor, confirm prices, submit. After approval, click Send to commit the order "
              "to the vendor. No accounting entry yet — a PO is a commitment, not a transaction.")

    add_h(doc, "Stage B — Goods Receipt (first JE)", 2)
    add_p(doc, "Operations → Goods Receipts → New from PO.")
    add_p(doc, "The receiving team records what physically arrived. Each line "
              "defaults to the remaining quantity from the PO but can be "
              "reduced for partial deliveries. Click Post on the Draft GR — "
              "this is the first accounting event:")
    add_bullet(doc, "GR transitions to Posted (no further edits).")
    add_bullet(doc, "Parent PO's open quantity decrements. If fully received, the PO auto-closes.")
    add_bullet(doc, "Finance creates a journal entry: DR the item's inventory "
                    "account / CR GR/IR Clearing.")
    add_p(doc, "After posting, a green banner appears with 'View Journal Entry "
              "→' linking to the new JE.")
    add_p(doc, "Concrete example: 10 bags of Tomato Seeds at AED 100/bag:")
    add_code(doc, "DR  121000-001 Raw Materials - Seeds          1,000.00\n"
                  "CR  223000-004 Goods Received Not Invoiced              1,000.00")

    # ───────────────────────────────────────────────────────────────────────
    add_h(doc, "Stage C — AP Invoice (second JE, with three-way match)", 2)
    add_p(doc, "When the vendor's invoice arrives, navigate to Operations → "
              "Purchasing → AP Invoices → New from GR. Pick the posted GR. "
              "The system pre-fills:")
    add_bullet(doc, "Each line's quantity is locked — matches what was received.")
    add_bullet(doc, "Each line's PO unit price is read-only.")
    add_bullet(doc, "Each line's invoice unit price defaults to the PO price "
                    "and is editable — adjust if the vendor charged a different amount.")
    add_p(doc, "Enter the header fields:")
    add_bullet(doc, "Vendor Invoice Number — as printed on the vendor's document.")
    add_bullet(doc, "Invoice Date — date on the vendor's document.")
    add_bullet(doc, "Due Date — defaults to invoice date + 30 days; override if "
                    "your payment terms differ.")
    add_bullet(doc, "Tax Code per line — defaults to S (Standard 5%). Pick SR "
                    "for reverse-charge imports of services.")
    add_p(doc, "Submit. The doc goes to Pending Approval per the AP_INVOICE "
              "approval rule (default: accountant approves over AED 10,000).")

    add_h(doc, "How three-way match works in v1", 3)
    add_p(doc, "The system compares three documents — the PO, the GR, and the "
              "invoice. In v1:")
    add_bullet(doc, "Quantity is forced to match the GR (no partial invoicing in v1).")
    add_bullet(doc, "Price may differ from the PO — the difference is a 'price "
                    "variance' computed per line and aggregated on the header.")
    add_bullet(doc, "Variance shows in the form as you edit prices, both per "
                    "line (red if positive, green if negative) and as a header total.")
    add_bullet(doc, "On approval, any non-zero variance posts to the Purchase "
                    "Price Variance account configured in Posting Setup.")

    add_h(doc, "The AP Invoice JE", 3)
    add_p(doc, "Example: same 10 bags of seeds, but vendor invoiced at AED "
              "105/bag (5 over the PO price) with 5% VAT:")
    add_code(doc, "DR  223000-004 Goods Received Not Invoiced     1,000.00   (= expected net, clears GR holding)\n"
                  "DR  514000-004 Purchase Price Variance             50.00   (= over-billed amount)\n"
                  "DR  122000-001 Input VAT Recoverable               52.50   (5% of 1,050 invoiced amount)\n"
                  "CR  221000-001 Trade Payables - Suppliers                1,102.50   (vendor's specific liability)")
    add_p(doc, "The GR's temporary holding (GR/IR Clearing) is now zero. The "
              "vendor liability is on the books. VAT is recorded as an asset "
              "reclaimable from the tax authority.")

    add_h(doc, "Reverse-charge VAT (UAE-specific)", 3)
    add_p(doc, "For imports of services from outside the UAE and "
              "designated-zone scenarios, the UAE FTA requires self-accounted "
              "VAT — the buyer records both Input VAT (claim back) and Output "
              "VAT (owe to tax authority) for the same amount.")
    add_p(doc, "On any AP Invoice line, pick the SR tax code instead of S to "
              "trigger reverse charge. The system automatically produces both "
              "VAT entries:")
    add_code(doc, "DR  223000-004 Goods Received Not Invoiced     1,000.00\n"
                  "DR  122000-001 Input VAT Recoverable               50.00\n"
                  "CR  222000-001 Output VAT Payable                              50.00\n"
                  "CR  221000-001 Trade Payables - Suppliers                  1,000.00")
    add_p(doc, "Note: the AP credit is the net amount (1,000) — not 1,050 — "
              "because the foreign vendor did not bill us UAE VAT. Net effect "
              "on the books is zero VAT cash, but both sides are recorded for "
              "VAT-return compliance.")

    add_h(doc, "Tax point (UAE Article 25)", 3)
    add_p(doc, "The system computes the VAT tax point as the earliest of: "
              "date of supply (= GR date), invoice date, payment date. The "
              "resolved date appears in the VAT line's description so it can "
              "be reconstructed for VAT returns.")

    # ───────────────────────────────────────────────────────────────────────
    add_h(doc, "Stage D — Vendor Payment (third JE)", 2)
    add_p(doc, "Navigate to Finance → Vendor Payments → New Payment.")
    add_p(doc, "Step 1: pick the vendor. The form will show that vendor's open "
              "AP invoices (Approved, not fully paid).")
    add_p(doc, "Step 2: pick the bank account (defaults to the one configured "
              "in Posting Setup), payment method (Bank Transfer / Cheque / "
              "Cash), payment date, and reference number (required when method "
              "= cheque — cheque number).")
    add_p(doc, "Step 3: check the invoices to pay. The 'Amount to Apply' input "
              "for each invoice pre-fills with the full outstanding amount — "
              "reduce it for a partial payment. The running total at the "
              "bottom shows what's being paid in this transaction.")
    add_p(doc, "Submit. The system creates the payment record and the JE in "
              "one atomic action:")
    add_code(doc, "DR  221000-001 Trade Payables - Suppliers      1,050.00\n"
                  "CR  126000-002 Cash at Bank - AED Operating              1,050.00")
    add_p(doc, "After this, AP Control is back to zero for this vendor. The "
              "payment is immutable — to correct a mistake, reverse the JE "
              "(Phase D.5 below).")

    # ───────────────────────────────────────────────────────────────────────
    add_h(doc, "Part 3 — Reading the Books", 1)

    add_h(doc, "Journal Entries", 2)
    add_p(doc, "Finance → Journal Entries. Read-only register of every posted "
              "JE. Click a row to expand and see the individual debit/credit "
              "lines with account names. Filter by source event type, status, "
              "date range. Voided JEs are shown struck-through; reversal JEs "
              "carry a 'Reversal' badge.")

    add_h(doc, "Trial Balance", 2)
    add_p(doc, "Finance → Trial Balance. The foundational accountant report. "
              "Aggregates every account's debit and credit total as of a "
              "chosen date and proves that total debits equal total credits. "
              "Filter by company, period, or include voided JEs.")
    add_p(doc, "If totals don't equal (which should never happen), the report "
              "highlights this in red — call your administrator immediately.")

    add_h(doc, "Vendor Sub-Ledger", 2)
    add_p(doc, "Finance → Vendor Sub-Ledger. Shows the per-vendor breakdown "
              "of the AP Control account balance. After the demo cycle above, "
              "the vendor's balance should be zero (we owe them, we paid them).")

    add_h(doc, "AP Aging", 2)
    add_p(doc, "Finance → AP Aging. Shows outstanding AP grouped by how "
              "overdue each invoice is: Not Due / 1-30 / 31-60 / 61-90 / >90 "
              "days. Per-vendor breakdown sorted by total amount descending. "
              "Critical for cash-flow planning and vendor relationships.")

    # ───────────────────────────────────────────────────────────────────────
    add_h(doc, "Part 4 — Corrections and Period Control", 1)

    add_h(doc, "Reversing a wrong journal entry", 2)
    add_p(doc, "Open the JE detail. Click 'Reverse Entry'. Provide a reason "
              "(required, 5-500 chars). The system:")
    add_bullet(doc, "Marks the original JE as 'Voided' (not deleted — audit trail preserved).")
    add_bullet(doc, "Creates a new JE with debits and credits swapped, dated today.")
    add_bullet(doc, "Links the two via sourceDocId for traceability.")
    add_p(doc, "Restricted to finance_admin, admin, super_admin. The reversal "
              "posts in today's open period — you cannot back-date corrections "
              "into closed periods.")

    add_h(doc, "Closing a fiscal period", 2)
    add_p(doc, "Finance → Fiscal Periods. Find the period to close, click "
              "'Close Period', provide a reason. After close:")
    add_bullet(doc, "No new postings can land in that period (GR, AP Invoice, "
                    "Payment will all fail with 'No open fiscal period'.")
    add_bullet(doc, "Existing JEs in the period are unaffected.")
    add_bullet(doc, "An audit trail records who closed it, when, and why.")
    add_p(doc, "To allow back-posting after close, reopen the period (also "
              "audited; reason required). Close it again afterwards.")

    # ───────────────────────────────────────────────────────────────────────
    add_h(doc, "Glossary", 1)
    add_table(doc,
        headers=["Term", "Meaning"],
        rows=[
            ["GL / General Ledger", "Master record of every accounting entry. Made up of accounts grouped into drawers."],
            ["Chart of Accounts (CoA)", "The list of accounts that make up your GL."],
            ["Journal Entry (JE)", "One balanced transaction made of 2+ lines. Total DR always equals total CR."],
            ["Debit (DR) / Credit (CR)", "The two sides of every entry. Sign depends on account type — assets/expenses increase on DR; liabilities/equity/revenue increase on CR."],
            ["PR / PO / GR / AP Invoice / Payment", "The five documents of the procure-to-pay cycle."],
            ["GR/IR Clearing", "Temporary accrued liability between goods receipt and vendor invoice."],
            ["AP / Accounts Payable", "What we owe vendors. A liability."],
            ["VAT", "Value-Added Tax. UAE standard rate 5%. Input VAT (purchases) is reclaimable; Output VAT (sales) is payable."],
            ["Reverse-charge VAT", "UAE requirement on imports of services — buyer self-accounts both Input and Output VAT for the same amount."],
            ["Tax point", "The legal date on which VAT is recognised. UAE FTA Article 25: earliest of supply / invoice / payment date."],
            ["Fiscal Period", "An accounting time-window, usually a month. Postings can only land in open periods."],
            ["Control Account", "A GL account managed by the system (Trade Payables, Trade Receivables). Not user-editable."],
            ["Three-way Match", "The check at vendor-invoice time that the invoice agrees with the PO and the GR."],
            ["Price Variance", "Difference between PO price and invoice price — posts to a dedicated variance account."],
            ["Sub-Ledger", "Per-counterparty detail behind a Control Account total. Sums always equal the Control balance."],
            ["Reversal JE", "An offsetting journal entry that voids an earlier one. Required for corrections — entries are never edited or deleted."],
        ])

    add_h(doc, "Support", 1)
    add_p(doc, "If a page reports an error or behaves unexpectedly, capture "
              "the page URL, the error text, the time, and which JE/doc you "
              "were viewing. Forward to your platform administrator.")
    add_p(doc, "Additional reference documents are under Docs/4-Finance-Mod-docs/ in the project repository.")

    return doc


# ============================================================================
# 2. PHASE C + D PM REPORT
# ============================================================================

def build_phase_c_d_pm_report():
    doc = Document()
    style_doc(doc)

    t = doc.add_heading("A64 Core Platform — Phase C + D Status Report", level=0)
    for r in t.runs:
        r.font.color.rgb = RGBColor(0x1F, 0x36, 0x56)
    sub = doc.add_paragraph()
    sub.add_run("Vendor Invoicing, Three-Way Match, Payment, and Closing the P2P Cycle").italic = True
    meta = doc.add_paragraph()
    meta.add_run("Date: 2026-05-21    ·    Author: Viet Anh").italic = True

    add_h(doc, "Executive Summary", 1)
    add_p(doc,
        "The procure-to-pay cycle is now closed. A purchase request can be "
        "raised, approved, converted to a purchase order, received as goods, "
        "invoiced by the vendor, and paid — with three journal entries posting "
        "automatically as the documents move through their lifecycle. The full "
        "cycle is demonstrable end-to-end in the development environment with "
        "no manual posting required.")
    add_p(doc,
        "Phase C (AP Invoice + three-way match) and Phase D (Payment) shipped "
        "in this period. Phase D.5 (Period Close + audit lock) was added per "
        "PM feedback to allow controllers to lock prior periods immediately "
        "rather than waiting for a later phase. The six critical items in the "
        "PM's 'Mistakes / risks from the mock system' note were also resolved "
        "in-flight (see separate FEEDBACK_RESPONSE.docx).")

    add_h(doc, "What Now Works End-to-End", 1)
    add_p(doc, "The platform produces three real journal entries per "
              "procurement cycle. Concrete example for 10 bags of fertilizer "
              "at AED 105/bag (vendor invoiced 5 above the PO price of 100), "
              "5% VAT:")
    add_table(doc,
        headers=["Stage", "Journal entry produced"],
        rows=[
            ["Goods Receipt", "DR Inventory 1,000.00 / CR GR/IR Clearing 1,000.00"],
            ["AP Invoice", "DR GR/IR Clearing 1,000.00 + DR Purchase Price Variance 50.00 + DR Input VAT 52.50 / CR AP - Vendor 1,102.50"],
            ["Vendor Payment", "DR AP - Vendor 1,102.50 / CR Cash at Bank 1,102.50"],
        ])
    add_p(doc, "After the cycle: Inventory holds the asset. The vendor liability is cleared. VAT is recorded as an asset reclaimable from the FTA. Bank balance reduced. Books balance.")

    add_h(doc, "What Was Built in Phase C — AP Invoice and Three-Way Match", 1)
    add_table(doc,
        headers=["Capability", "Description"],
        rows=[
            ["AP Invoice document", "New operational document type AP, created from a posted GR. Lifecycle: Draft → Pending Approval → Approved (terminal) or Rejected (terminal)."],
            ["Three-way match logic", "Quantity locked from the GR (no partial invoicing in v1). Price editable. Per-line and header variance computed and displayed in real time."],
            ["Approval flow", "Reuses the approval engine. Default rule: AP_INVOICE over AED 10,000 requires accountant approval. Configurable in Approval Rules."],
            ["AP Invoice list / detail / form pages", "Operations → Purchasing → AP Invoices. Variance UI shows positive variance in red, negative in green. Cross-page link to view the resulting JE."],
            ["Posting handler", "Finance side. JE produces DR GR/IR Clearing + DR Input VAT + DR/CR Price Variance + CR AP - Vendor. Sum balances. AP credit aggregates per-line based on tax code (standard vs reverse charge)."],
            ["Purchase Price Variance account", "514000-004, seeded into the CoA. Posting Setup field populated."],
            ["Reverse-charge VAT support", "Per-line tax code SR triggers dual VAT entries (DR Input + CR Output, same amount). Required for UAE imports of services. New isReverseCharge flag on tax_codes table."],
            ["UAE Article 25 tax point", "VAT tax point computed as earliest of date-of-supply (GR date) and invoice date. Recorded on the VAT line description for VAT-return reconstruction."],
        ])

    add_h(doc, "What Was Built in Phase D — Vendor Payment", 1)
    add_table(doc,
        headers=["Capability", "Description"],
        rows=[
            ["Vendor Payment module", "Finance-side action (not operation). Finance team picks open AP invoices and records a payment with method, reference, bank account."],
            ["Open-invoice picker", "Per-vendor list of Approved AP invoices with outstanding amount = totalGross minus prior payment applications. Multi-invoice payment supported on a single transaction."],
            ["Payment number generation", "Format PAY-{companyCode}-{YYYY}-{NNNN}, monotonic, unique."],
            ["Atomic JE creation", "Payment record and JE post in the same database transaction. JE = DR AP Control + CR Bank, balanced."],
            ["Payment list / detail pages", "Finance → Vendor Payments. List with filters by vendor, date range, search. Detail page links to the applied invoices and the underlying JE."],
            ["Cross-store join (open invoices)", "Frontend orchestrates: fetches open AP from operation, joins with totalsPaid from finance, computes outstanding client-side. v1 trade-off; avoids new service-to-service runtime call."],
        ])

    add_h(doc, "What Was Built in Phase D.5 — Period Close + Audit Lock", 1)
    add_p(doc, "Promoted from Phase E.2 per PM feedback. The audit lock is "
              "now enforceable from day one.")
    add_table(doc,
        headers=["Capability", "Description"],
        rows=[
            ["Fiscal Periods management UI", "Finance → Fiscal Periods. List with status pills (Open/Closed), Current-period badge, close/reopen actions."],
            ["Bulk-create wizard", "Create a full fiscal year of periods (12 monthly, 13 for 4-4-5, or 4 quarterly) in one action. Configurable initial status (all open vs all closed except current)."],
            ["Close-period enforcement", "Every posting handler resolves the fiscal period and refuses to post if the period is closed. Affects GR, AP Invoice, Payment, and JE reversal."],
            ["Audit trail", "Close and reopen actions are audited — closedBy, closedAt, closeReason and the corresponding reopened fields are persisted and surfaced."],
        ])

    add_h(doc, "PM Feedback Items Addressed In-Flight", 1)
    add_p(doc, "Six critical items from the PM's 'Mistakes / risks from the mock system' note were resolved in the same period as Phases C and D. Reference: FEEDBACK_RESPONSE.docx.")
    add_table(doc,
        headers=["#", "Issue", "Status"],
        rows=[
            ["1", "GR/IR account misclassified as Trade Payable", "Reclassified into a new 223000 Accruals & Deferred Income group; old account deactivated; existing Posting Setup auto-migrated."],
            ["4", "No JE reversal mechanism", "Reverse Entry action implemented. Original status flips to void; offsetting JE created with debits/credits swapped. Reason required, audited."],
            ["5", "No Trial Balance report", "New /finance/trial-balance page and endpoint. Aggregates all JE lines per account, validates total DR equals total CR."],
            ["6", "Period close deferred too late", "Promoted from Phase E to Phase D.5 and shipped this period."],
            ["10", "No Rounding Differences account in seed", "617000-011 added to default CoA."],
            ["11", "Per-item valuation method violates IAS 2", "Moved from per-item to company-level (Posting Setup → Inventory Valuation section)."],
            ["12", "Purchase Price Variance missing from seed", "514000-004 added to default CoA."],
        ])

    add_h(doc, "PM Feedback Items Scheduled for Later", 1)
    add_table(doc,
        headers=["#", "Issue", "Where"],
        rows=[
            ["2", "UAE Article 25 tax-point rule not applied", "Shipped in this period along with Phase C"],
            ["3", "No reverse-charge VAT mechanism", "Shipped in this period along with Phase C"],
            ["7", "No multi-company UI", "Tracked; not yet scheduled"],
            ["8", "No Biological Assets fair-value P&L line (IAS 41)", "Phase E"],
            ["9", "Direct Labour EOSB in OpEx instead of COGS (IAS 2)", "Phase E"],
            ["13", "Native dropdown jumps to selected on reopen", "Phase E (UI polish)"],
        ])
    add_p(doc, "Items 2 and 3 originally tagged for Phase D were brought "
              "forward and shipped as refinements alongside Phase C work this "
              "period. The remaining items are non-blocking for the v1 P2P cycle.")

    add_h(doc, "Build Metrics", 1)
    add_p(doc, "Engineering output during this period (Phases C + D + D.5 + PM refinements):")
    add_bullet(doc, "7 new finance Alembic migrations (007 through 013).")
    add_bullet(doc, "Operation backend: 1 new doc type (AP), 14 endpoints under /api/v1/purchasing/ap/.")
    add_bullet(doc, "Finance backend: 2 new posting handlers, 3 new reports (Trial Balance, AP Aging, Vendor Sub-Ledger), 1 reversal action, 4 new endpoints under /api/v1/finance/reports/, fiscal periods audit trail, payment recording.")
    add_bullet(doc, "Frontend: 9 new pages (AP Invoice list/detail/form, Vendor Payment list/detail/form, Trial Balance, AP Aging, Vendor Sub-Ledger, Fiscal Periods).")
    add_bullet(doc, "Tests: 100+ new test cases across backend modules. All in CI.")
    add_bullet(doc, "Verified: complete PR → PO → GR → AP Invoice → Payment cycle end-to-end with three real journal entries posting against a real chart of accounts.")

    add_h(doc, "Remaining Work to Reach v1 Production", 1)
    add_p(doc, "Outside of the operational P2P cycle, the following is still pending before customer rollout:")
    add_bullet(doc, "Production deployment plan: MySQL hosting, secrets, backups, monitoring, finance-stack docker-compose.prod.yml.")
    add_bullet(doc, "Sales-side / AR module. We have built P2P only. Customer invoicing and receivables are the other half of every accounting system.")
    add_bullet(doc, "GR/IR reconciliation report (Phase E.1): month-end check that the clearing account zeros out.")
    add_bullet(doc, "Customer onboarding playbook documenting the day-one setup steps in this guide.")
    add_bullet(doc, "Multi-company UI — the data model supports it; the admin screen is not yet built (PM item 7).")
    add_bullet(doc, "Phase E polish items: biological-assets accounting, direct-labour EOSB split, UX consistency on native dropdowns.")

    add_h(doc, "Controls Debt to Resolve Before Production", 1)
    add_p(doc, "Three dev-environment conveniences relax controls that real accounting policy normally enforces. Each must be decided before customer rollout:")
    add_table(doc,
        headers=["Item", "Decision required"],
        rows=[
            ["admin / super_admin can approve their own documents", "Gate behind a feature flag that defaults off in production, or remove and require separate user accounts for requesters vs approvers. Separation of duties is non-negotiable for auditor-tested compliance."],
            ["Posted GR cannot be reversed", "Add a reversal-GR flow before customer rollout if damaged-goods or wrong-receipt scenarios are expected."],
            ["Tax codes not validated server-side on PR/PO/AP create", "Frontend dropdown is the only enforcement. Add backend validation when the AP Invoice handler reads tax codes."],
        ])

    add_h(doc, "Risks and Open Questions", 1)
    add_table(doc,
        headers=["Risk", "Note"],
        rows=[
            ["Frontend-orchestrated cross-store joins (AP Aging, open-AP-payment picker)", "v1 trade-off chosen for simplicity. If volumes grow significantly, move to backend service-to-service joins."],
            ["Concurrent JE-number generation under high load", "Currently uses MAX+1 plus a UNIQUE constraint. Adequate for current scale; replace with dedicated counter table if contention becomes measurable."],
            ["Period close after Posted JE in that period", "Allowed but not blocked. Standard accounting practice. Reopen is audited."],
            ["No tax-rate version history", "Tax rate is stored as a single value per code. If UAE rates change, existing JEs keep the old rate (correct) but future calls use the new rate (also correct). No historical reporting issue."],
        ])

    add_h(doc, "What This Period Enables for the Customer", 1)
    add_p(doc, "After this period's work, the platform is capable of running a "
              "full UAE procurement cycle with auto-generated journal entries "
              "and a working set of finance reports. Specifically the "
              "customer can now:")
    add_bullet(doc, "Raise and approve a purchase request and order.")
    add_bullet(doc, "Receive goods from the vendor and post the receipt JE.")
    add_bullet(doc, "Record the vendor's invoice with three-way match and price-variance handling.")
    add_bullet(doc, "Handle UAE reverse-charge VAT correctly.")
    add_bullet(doc, "Pay the vendor and close out the AP balance.")
    add_bullet(doc, "Run the Trial Balance, AP Aging, and Vendor Sub-Ledger reports for accounting review.")
    add_bullet(doc, "Open and close fiscal periods with full audit trail.")
    add_bullet(doc, "Reverse a wrong journal entry with proper audit linkage.")
    add_p(doc, "This is the minimum feature set to demonstrate a working "
              "accounting system to a finance audience. AR (sales-side) is "
              "the next major build to round out the v1 to a complete ERP.")

    add_h(doc, "Reference", 1)
    add_bullet(doc, "POSTING_ENGINE_ROADMAP.md — phased build plan + controls-debt register.")
    add_bullet(doc, "INTEGRATION_MODEL.md — cross-service event design and immutability rules.")
    add_bullet(doc, "FEEDBACK_RESPONSE.docx — itemised response to the PM's feedback note.")
    add_bullet(doc, "FINANCE_USER_GUIDE.docx — end-user walkthrough of the full P2P cycle.")
    add_bullet(doc, "FINANCE_MODULE_GUIDE.md — living architecture document.")

    return doc


def main():
    user_doc = build_user_guide()
    user_path = OUT / "FINANCE_USER_GUIDE.docx"
    user_doc.save(str(user_path))
    print(f"wrote {user_path}")

    pm_doc = build_phase_c_d_pm_report()
    pm_path = OUT / "FINANCE_PM_REPORT_PHASE_C_D.docx"
    pm_doc.save(str(pm_path))
    print(f"wrote {pm_path}")


if __name__ == "__main__":
    main()
