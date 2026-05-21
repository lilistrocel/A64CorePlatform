"""
Generate FEEDBACK_RESPONSE.docx — itemised reply to the PM's
"Mistakes / risks from the mock system" feedback note.

Output: Docs/4-Finance-Mod-docs/FEEDBACK_RESPONSE.docx
"""

from pathlib import Path

from docx import Document
from docx.shared import Cm, Pt, RGBColor

OUT = Path(__file__).resolve().parent.parent / "Docs" / "4-Finance-Mod-docs" / "FEEDBACK_RESPONSE.docx"


def style_doc(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25


def add_h(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x36, 0x56)


def add_p(doc, text, italic=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.italic = italic


def add_bullet(doc, text):
    p = doc.add_paragraph(text, style="List Bullet")
    p.paragraph_format.left_indent = Cm(0.75)


def add_table(doc, headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Light Grid Accent 1"
    for i, h in enumerate(headers):
        t.rows[0].cells[i].text = h
        for run in t.rows[0].cells[i].paragraphs[0].runs:
            run.bold = True
    for ri, row in enumerate(rows, start=1):
        for ci, val in enumerate(row):
            t.rows[ri].cells[ci].text = val
            for p in t.rows[ri].cells[ci].paragraphs:
                p.paragraph_format.space_after = Pt(2)
    doc.add_paragraph()


def issue_section(doc, num, title, raised, fix, status):
    add_h(doc, f"Item {num}. {title}", 2)

    p = doc.add_paragraph()
    p.add_run("Status: ").bold = True
    color_run = p.add_run(status)
    color_run.bold = True
    if "Shipped" in status:
        color_run.font.color.rgb = RGBColor(0x10, 0x82, 0x37)
    elif "Phase D" in status or "Phase D.5" in status:
        color_run.font.color.rgb = RGBColor(0xB4, 0x53, 0x09)
    elif "Phase E" in status:
        color_run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
    else:
        color_run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

    add_p(doc, "")
    p = doc.add_paragraph()
    p.add_run("Raised: ").bold = True
    p.add_run(raised)

    p = doc.add_paragraph()
    p.add_run("Response: ").bold = True
    p.add_run(fix)
    doc.add_paragraph()


def main():
    doc = Document()
    style_doc(doc)

    # Title
    t = doc.add_heading("Response to PM Feedback — Mock System Review", level=0)
    for r in t.runs:
        r.font.color.rgb = RGBColor(0x1F, 0x36, 0x56)
    sub = doc.add_paragraph()
    sub.add_run("Itemised response to the PM's feedback note ").italic = True
    sub.add_run("\"Mistakes / risks from the mock system\".").italic = True
    meta = doc.add_paragraph()
    meta.add_run("Date: 2026-05-20    ·    Author: Viet Anh").italic = True

    add_h(doc, "Summary", 1)
    add_p(doc,
        "Thirteen issues were raised in the feedback note. Six were classified as "
        "critical (compliance-grade or fundamental functionality gaps) and were "
        "resolved in the same session before continuing to Phase D. Three are "
        "scheduled into Phase D as part of the planned VAT and period-close work. "
        "Three more are scheduled into Phase E as part of the operability and "
        "polish work. One (multi-company UI) is tracked but not yet scheduled."
    )

    add_h(doc, "Disposition table", 1)
    add_table(doc,
        headers=["#", "Issue (one-line)", "Disposition"],
        rows=[
            ["1", "GR/IR classified as Trade Payable", "Shipped 2026-05-20"],
            ["2", "Input VAT tax point rule (UAE Article 25)", "Phase D"],
            ["3", "No reverse-charge VAT mechanism", "Phase D"],
            ["4", "No JE reversal mechanism", "Shipped 2026-05-20"],
            ["5", "No Trial Balance report", "Shipped 2026-05-20"],
            ["6", "Period close deferred too late", "Promoted to Phase D.5"],
            ["7", "No multi-company UI", "Tracked, no phase yet"],
            ["8", "No Biological Assets fair-value P&L line (IAS 41)", "Phase E"],
            ["9", "Direct Labour EOSB in OpEx not COGS (IAS 2)", "Phase E"],
            ["10", "No Rounding Differences account", "Shipped 2026-05-20"],
            ["11", "Per-item valuation method violates IAS 2", "Shipped 2026-05-20"],
            ["12", "Purchase Price Variance missing from seed", "Shipped 2026-05-20"],
            ["13", "Native dropdown jumps to selected on reopen", "Phase E"],
        ])

    add_h(doc, "Shipped in this session (six items)", 1)

    issue_section(doc, 1,
        "GR/IR classified as Trade Payable",
        raised=(
            "The Goods Received Not Invoiced account (221000-002) was placed under "
            "the Trade Payables (221000) header. GR/IR is an accrued liability — "
            "goods have arrived but no formal vendor invoice exists yet. Putting it "
            "under Trade Payables corrupts the AP aging report and the sub-ledger "
            "reconciliation."
        ),
        fix=(
            "Created a new accrued-liabilities section 223000 Accruals & Deferred "
            "Income and added 223000-004 Goods Received Not Invoiced under it. "
            "The old 221000-002 account was deactivated (isActive=false) — not "
            "deleted, because historical journal entries reference it and "
            "deletion would break audit traceability. The existing Posting Setup "
            "row was automatically migrated to point at the new account. New "
            "deployments pick up the correct structure via the updated seed loader "
            "(default_coa.py + seed_loader.py)."
        ),
        status="Shipped 2026-05-20",
    )

    issue_section(doc, 4,
        "No JE reversal mechanism",
        raised=(
            "Standard accounting requires reversing a wrong journal entry with an "
            "offsetting JE that references the original — never delete or edit. "
            "The system had no such action. Schema columns existed (status='void', "
            "voidedBy, voidedAt, voidReason) but no API to flip them."
        ),
        fix=(
            "New endpoint POST /api/v1/finance/journal-entries/{jeId}/reverse + a "
            "Reverse Entry button on the JE detail view. The action: (1) loads the "
            "original JE; (2) refuses if already void; (3) ensures a current open "
            "fiscal period exists; (4) inserts a new JE with debits and credits "
            "swapped, with sourceDocId pointing back to the original; (5) flips the "
            "original to void with voidedBy/voidedAt/voidReason recorded. Reason is "
            "required, 5-500 chars. Restricted to finance_admin / admin / "
            "super_admin. The reversal posts in today's date so it cannot be used "
            "to back-date corrections into closed periods."
        ),
        status="Shipped 2026-05-20",
    )

    issue_section(doc, 5,
        "No Trial Balance report",
        raised=(
            "A trial balance is the foundational accountant view — every account's "
            "net debit/credit balance as of a date. Auditors and accountants ask "
            "for it first. The system had no such report."
        ),
        fix=(
            "New endpoint GET /api/v1/finance/reports/trial-balance + a Trial "
            "Balance page in the Finance sidebar. Aggregates journal_entry_lines "
            "by account with filters for company, as-of date, period, and include-"
            "voided. Includes accounts with zero activity so accountants see the "
            "full CoA. Footer validates total debits equal total credits — if not, "
            "the page surfaces a red 'books out of balance' warning. SQL uses a "
            "subquery aggregation to handle the LEFT JOIN edge cases correctly."
        ),
        status="Shipped 2026-05-20",
    )

    issue_section(doc, 10,
        "No Rounding Differences account in the seed CoA",
        raised=(
            "The Posting Setup page exposes a Rounding Differences field, but the "
            "seed CoA had no matching account for users to point it at."
        ),
        fix=(
            "Added 617000-011 Rounding Differences (Operating Cost · General & "
            "Administrative) to the seed loader and inserted into the dev DB."
        ),
        status="Shipped 2026-05-20",
    )

    issue_section(doc, 11,
        "Per-item valuation method violates IAS 2 consistency",
        raised=(
            "IAS 2 requires the same cost formula across inventories of similar "
            "nature and use. The Item GL Mapping page exposed Moving Average / "
            "Standard / FIFO per item — users could pick FIFO on one seed and "
            "Moving Average on another. That is an audit finding."
        ),
        fix=(
            "Removed the Valuation Method column from the Item Mapping page. Added "
            "a new Inventory Valuation section to the Posting Setup page with a "
            "single company-level choice (defaultValuationMethod). The field "
            "moved from purchase_item_finance_ext to company_posting_setup. The "
            "per-item column is preserved in the data model for backward "
            "compatibility but is no longer editable and no longer authoritative."
        ),
        status="Shipped 2026-05-20",
    )

    issue_section(doc, 12,
        "Purchase Price Variance account missing from seed",
        raised=(
            "Phase C posting handler needs a Purchase Price Variance GL account "
            "to absorb the difference when vendor invoice price differs from PO "
            "price. The seed CoA had no such account. The dev doc had noted "
            "'leave PPV blank in v1, add in Phase C' — but the account itself "
            "didn't exist anywhere yet."
        ),
        fix=(
            "Added 514000-004 Purchase Price Variance (Cost of Sales · Inventory "
            "Adjustments) to the seed loader. Already inserted into the dev DB "
            "ahead of Phase C, so the AP Invoice posting handler can resolve to "
            "it correctly. Posting Setup field purchasePriceVarianceAccountId can "
            "now be populated with this account."
        ),
        status="Shipped 2026-05-20",
    )

    add_h(doc, "Scheduled into Phase D (three items)", 1)

    issue_section(doc, 2,
        "Input VAT date does not follow UAE Article 25",
        raised=(
            "UAE VAT Decree-Law Article 25 defines the tax point as the earliest "
            "of: date of supply, date of invoice, date of payment, or partial-"
            "payment date. The current AP Invoice handler uses the invoice date "
            "alone."
        ),
        fix=(
            "Phase D will carry both the date of supply (GR date) and the invoice "
            "date through the contract. The posting handler will resolve the tax "
            "point as min(grDate, invoiceDate, paymentDate-if-any) and store the "
            "resolved value on the JE for VAT-return reporting."
        ),
        status="Phase D",
    )

    issue_section(doc, 3,
        "No reverse-charge VAT mechanism",
        raised=(
            "Reverse charge is mandatory for UAE imports of services and "
            "designated-zone scenarios. The buyer self-accounts: DR Input VAT + "
            "CR Output VAT, same amount, same line. The SR tax code exists in the "
            "seed but the posting handler treats it like normal VAT — only a DR "
            "Input VAT line is created."
        ),
        fix=(
            "Phase D will: (a) add an isReverseCharge flag to the tax_codes "
            "table; (b) extend the AP Invoice handler to detect reverse-charge "
            "lines; (c) post a balancing CR Output VAT for the same amount in "
            "addition to the standard DR Input VAT. Net effect on the books is "
            "zero VAT cash, but both sides are recorded for compliance."
        ),
        status="Phase D",
    )

    issue_section(doc, 6,
        "Period close deferred too late",
        raised=(
            "Period close + audit lock was scheduled in Phase E. The PM noted "
            "audit lock is a control and should not be delayed."
        ),
        fix=(
            "Promoted to Phase D.5 (immediately after Payment in Phase D). The "
            "_resolve_fiscal_period_or_raise helper will refuse postings to closed "
            "periods. The UI to open/close periods uses the data model and enum "
            "that already exist."
        ),
        status="Phase D.5",
    )

    add_h(doc, "Scheduled into Phase E (three items)", 1)

    issue_section(doc, 8,
        "No Biological Assets fair-value remeasurement line",
        raised=(
            "IAS 41 requires biological assets (plants, livestock) to be measured "
            "at fair value less costs to sell, with the gain or loss going through "
            "P&L. The CoA has no Gain or Loss on Biological Assets account."
        ),
        fix=(
            "Phase E adds 812000-001 Gain on Biological Assets - Fair Value "
            "(Other Income) and 514000-005 Loss on Biological Assets - Fair Value "
            "(Cost of Sales) to the seed CoA. A new remeasurement event handler "
            "will post the fair-value adjustment. Domain-specific to agri."
        ),
        status="Phase E",
    )

    issue_section(doc, 9,
        "Direct Labour EOSB in OpEx instead of COGS",
        raised=(
            "Labour Overtime and Labour Benefits sit in 512000 Direct Labour, but "
            "611000-004 End of Service Benefits Expense is under Operating Cost. "
            "IAS 2 says direct production labour's full cost — including EOSB "
            "accrual — must be capitalised into inventory and flow through COGS, "
            "not expensed as OpEx."
        ),
        fix=(
            "Phase E adds 512000-006 Direct Labour - EOSB under the existing 512000 "
            "Direct Labour header. Indirect labour EOSB stays in 611000-004. The "
            "split allows production labour EOSB to capitalise into inventory "
            "consistent with IAS 2."
        ),
        status="Phase E",
    )

    issue_section(doc, 13,
        "Native dropdown jumps to selected position on reopen",
        raised=(
            "When a user selects a value from a native HTML select, then reopens "
            "it, the browser scrolls the option list so the selected item is in "
            "the middle of the visible window. To see other options the user has "
            "to scroll up. The PM asked this be consistent — always open from the "
            "top of the list."
        ),
        fix=(
            "The AccountCombobox component (used on Posting Setup and Item GL "
            "Mapping) already opens from the top because it is a custom typeahead "
            "rather than a native select. Phase E replaces remaining native "
            "select controls (tax codes, payment terms, company codes) with the "
            "same searchable combobox pattern so behaviour is uniform across the "
            "app."
        ),
        status="Phase E",
    )

    add_h(doc, "Tracked but not yet scheduled (one item)", 1)
    issue_section(doc, 7,
        "No multi-company UI",
        raised=(
            "The PM noted there is no UI for managing multiple companies, and "
            "asked whether the architecture supports it in future steps."
        ),
        fix=(
            "The data model is multi-company throughout — every finance table "
            "carries companyCode and organizationId. Master data, posting setup, "
            "and journal entries are all per-company. The seed creates one "
            "company (1000) for the demo. A multi-company UI is not yet scheduled "
            "into a phase; it will be prioritised when the first multi-entity "
            "customer is onboarded. No data-model changes will be required at "
            "that point."
        ),
        status="Tracked",
    )

    add_h(doc, "Notes on the disposition", 1)
    add_p(doc,
        "All six critical fixes were dispatched and verified end-to-end in the "
        "same session as the feedback was received. The roadmap "
        "(POSTING_ENGINE_ROADMAP.md, section 4) has been updated to reflect the "
        "new phase ordering — in particular Period Close was promoted from "
        "Phase E.2 to Phase D.5, immediately after Payment lands. The Posting "
        "Engine continues toward Phase D (vendor payment + JE) as the next "
        "scheduled phase.")

    doc.save(str(OUT))
    print(f"wrote {OUT}")


if __name__ == "__main__":
    main()
